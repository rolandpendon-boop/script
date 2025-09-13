from PIL import Image
import pytesseract
from docx import Document
import os

# Path to tesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# 📂 Folder containing images
input_folder = r"C:\Users\Boyet\Pictures\flipbooks"
output_folder = r"C:\Users\Boyet\Pictures\flipbooks\word_files"

# Create output folder if not exists
os.makedirs(output_folder, exist_ok=True)

# Loop through images in folder
for file in sorted(os.listdir(input_folder)):
    if file.lower().endswith((".png", ".jpg", ".jpeg")):
        img_path = os.path.join(input_folder, file)

        # Step 1: OCR (English text from image)
        text = pytesseract.image_to_string(Image.open(img_path))

        if text.strip():
            # Step 2: Create a new Word doc per image
            doc = Document()
            doc.add_heading(f"Text from {file}", level=2)
            doc.add_paragraph(text)

            # Step 3: Save as separate Word file
            output_file = os.path.join(output_folder, f"{os.path.splitext(file)[0]}.docx")
            doc.save(output_file)
            print(f"✅ Saved: {output_file}")

print("🎉 All images processed into separate Word files!")
