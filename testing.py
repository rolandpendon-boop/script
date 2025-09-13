import pytesseract
from PIL import Image
from docx import Document
import os

# 👇 Point to your Tesseract installation
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


folder = r"C:\Users\Boyet\Pictures\flipbooks"

folder2 = r"C:\Users\Boyet\Pictures\flipbooks\pdf"
for file in os.listdir(folder):
    if file.lower().endswith((".png", ".jpg", ".jpeg")):
        img_path = os.path.join(folder, file)
        text = pytesseract.image_to_string(Image.open(img_path))

        # Create new Word document for each image
        doc = Document()
        doc.add_paragraph(f"--- OCR from {file} ---")
        doc.add_paragraph(text)

        # Build output filename with .docx extension
        base_name = os.path.splitext(file)[0]   # removes .jpg/.png
        output_file = os.path.join(folder2, base_name + ".docx")

        doc.save(output_file)
        print(f"✅ Created: {output_file}")
